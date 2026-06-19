#!/usr/bin/env python3
"""
minerador.py — Fase 1: Varredura mecânica + roteamento exclusivo multi-tópico

Cada seção de cada .docx é roteada para a RAW do tópico que ela melhor representa.
Uma seção nunca vai para dois tópicos — o roteamento é exclusivo.

Uso:
    python minerador.py                  # roda todos os tópicos
    python minerador.py --dry-run        # lista arquivos sem gravar nada
    python minerador.py --sem-embed      # só fuzzy (mais rápido, menos preciso)
    python minerador.py --debug-estilos  # inspeciona estilos do 1º docx e sai
    python minerador.py --reset          # apaga todos os RAWs e o manifesto (reprocessa tudo)
"""

import argparse
import json
import re
import sys
from datetime import datetime
from pathlib import Path

import docx
from rapidfuzz import fuzz
from tqdm import tqdm

# ─── Caminhos ─────────────────────────────────────────────────────────────────

RAIZ_SHAREPOINT = Path(
    r"C:\Users\Raphael.Caveagna\OneDrive - Invent"
    r"\Documentação _ Software - DOCUMENTAÇÃO - PASTA"
    r"\01 - WCS - DEVOPS (DOCUMENTAÇÃO)"
)

RAIZ_REDE = Path(r"R:\30.SOFTWARE-MFC\01 - Projetos")

RAIZ_COFRE    = Path(__file__).parent.parent / "corpus-conhecimento"
TOPICOS_JSON  = Path(__file__).parent / "topicos.json"
MANIFESTO_PATH = RAIZ_COFRE / "manifesto.json"
INBOX_PATH     = RAIZ_COFRE / "_INBOX_REVISAR.md"

# ─── Thresholds ───────────────────────────────────────────────────────────────

FUZZY_CONFIANTE  = 80    # score ≥ 80 → rota para o tópico vencedor direto
FUZZY_INCERTO    = 55    # 55–79 → checa com embedding
MARGEM_EXCLUSIVA = 18    # vencedor precisa superar o 2º colocado por ≥ 18 pts (fuzzy)
EMBED_CONFIANTE  = 0.65  # cosine ≥ 0.65 → RAW; abaixo → quarentena

# ─── Filtros de pasta ─────────────────────────────────────────────────────────

ANOS_VALIDOS = re.compile(r"^Especifica[çc][aã]o Funcional", re.IGNORECASE)

# ─── Helpers: estrutura de pastas ─────────────────────────────────────────────

def extrair_codigo_projeto(nome_pasta: str) -> str:
    m = re.match(r"^(I[\d.]+)", nome_pasta)
    return m.group(1) if m else nome_pasta.split(" - ")[0].strip()


# Nomes de arquivo a ignorar nos Aditivos (genéricos, sem conteúdo de processo)
_ADITIVOS_SKIP_NAMES = {
    "perguntas ptlsp.docx", "resposta ptlsp.docx", "analise.docx",
    "ajustes sistema velox.docx",
}
_ADITIVOS_SKIP_PREFIXES = ("pontos pendentes", "levantamento_funcional")

# Prefixos a ignorar na rede (templates e placeholders)
_REDE_SKIP_PREFIXES = ("ixx.", "template ", "(old)")


def _codigo_de_arquivo(nome: str) -> str:
    """Tenta extrair código Ixx.xxx do nome do arquivo (suporta I25_9043 também)."""
    m = re.search(r"I(\d{2})[._](\d+)", nome)
    if m:
        return f"I{m.group(1)}.{m.group(2)}"
    return ""


def descobrir_docx(raiz: Path, raiz_rede: Path | None = None) -> list[dict]:
    """Descobre .docx de Especificação Funcional, Aditivos e Rede da empresa."""
    encontrados = []

    # ── Especificação Funcional (estrutura: /ANO/PROJETO/*.docx) ────────────
    for pasta_ano in sorted(raiz.iterdir()):
        if not pasta_ano.is_dir() or not ANOS_VALIDOS.match(pasta_ano.name):
            continue
        ano = pasta_ano.name.split("-")[-1].strip()
        for pasta_proj in sorted(pasta_ano.iterdir()):
            if not pasta_proj.is_dir():
                continue
            codigo = extrair_codigo_projeto(pasta_proj.name)
            for f in pasta_proj.iterdir():
                if not f.is_file() or f.suffix.lower() != ".docx":
                    continue
                nome_lower = f.name.lower()
                if nome_lower.startswith("(old)"):
                    continue
                if "especifica" not in nome_lower and "especification" not in nome_lower:
                    continue
                encontrados.append({"caminho": f, "projeto": codigo, "ano": ano, "fonte": "ES"})

    # ── Aditivos (estrutura: /Aditivos/VENDOR/PROJETO/*.docx) ────────────────
    pasta_aditivos = raiz / "Aditivos"
    if pasta_aditivos.exists():
        for vendor in sorted(pasta_aditivos.iterdir()):      # FL-Soft, Velox
            if not vendor.is_dir():
                continue
            for subprojeto in sorted(vendor.iterdir()):      # DIAMANTE, BETA, etc.
                if not subprojeto.is_dir():
                    continue
                for f in subprojeto.iterdir():
                    if not f.is_file() or f.suffix.lower() != ".docx":
                        continue
                    nome_lower = f.name.lower()
                    if nome_lower in _ADITIVOS_SKIP_NAMES:
                        continue
                    if any(nome_lower.startswith(p) for p in _ADITIVOS_SKIP_PREFIXES):
                        continue
                    codigo = _codigo_de_arquivo(f.name) or subprojeto.name[:20]
                    encontrados.append({"caminho": f, "projeto": codigo, "ano": "ADITIVO", "fonte": "ADITIVO"})

    # ── Rede da empresa (estrutura: /YYYY/PROJETO/**/*.docx) ──────────────────
    if raiz_rede and raiz_rede.exists():
        for pasta_ano in sorted(raiz_rede.iterdir()):
            if not pasta_ano.is_dir() or not re.match(r"^\d{4}$", pasta_ano.name):
                continue
            ano = pasta_ano.name
            for pasta_proj in sorted(pasta_ano.iterdir()):
                if not pasta_proj.is_dir():
                    continue
                codigo = extrair_codigo_projeto(pasta_proj.name)
                for f in pasta_proj.rglob("*.docx"):   # recursivo — subpastas incluídas
                    if not f.is_file():
                        continue
                    nome_lower = f.name.lower()
                    if any(nome_lower.startswith(p) for p in _REDE_SKIP_PREFIXES):
                        continue
                    if "especifica" not in nome_lower and "especification" not in nome_lower:
                        continue
                    encontrados.append({"caminho": f, "projeto": codigo, "ano": ano, "fonte": "REDE"})

    return encontrados

# ─── Helpers: extração de seções ──────────────────────────────────────────────

_HEADING_PREFIXES = ("heading", "título", "titulo")

def _eh_heading(p) -> bool:
    return any(p.style.name.lower().startswith(pref) for pref in _HEADING_PREFIXES)


def extrair_secoes(caminho: Path) -> list[dict]:
    doc = docx.Document(str(caminho))
    secoes, heading_atual, buf = [], "[SEM TÍTULO]", []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if _eh_heading(p):
            if buf:
                secoes.append({"heading": heading_atual, "texto": "\n".join(buf)})
            heading_atual, buf = t, []
        else:
            buf.append(t)
    if buf:
        secoes.append({"heading": heading_atual, "texto": "\n".join(buf)})
    return secoes


def debug_estilos(raiz: Path):
    arquivos = descobrir_docx(raiz)
    if not arquivos:
        print("[ERRO] Nenhum .docx encontrado.")
        return
    caminho = arquivos[0]["caminho"]
    print(f"Arquivo: {caminho.name}\n")
    doc = docx.Document(str(caminho))
    vistos = {}
    for p in doc.paragraphs:
        s = p.style.name
        if s not in vistos:
            vistos[s] = p.text[:80]
    for nome, exemplo in vistos.items():
        print(f"  [{nome}]  ->  {exemplo}")

# ─── Helpers: vocabulário ─────────────────────────────────────────────────────

def carregar_vocabulario(path: Path) -> list[str]:
    termos = []
    for linha in path.read_text(encoding="utf-8").splitlines():
        linha = linha.strip()
        if linha and not linha.startswith("#"):
            termos.append(linha.lower())
    return termos

# ─── Helpers: match e roteamento ──────────────────────────────────────────────

def score_fuzzy(secao: dict, termos: list[str]) -> float:
    alvo = (secao["heading"] + " " + secao["texto"][:500]).lower()
    return max(fuzz.partial_ratio(t, alvo) for t in termos)


def score_semantico(secao: dict, embeddings, model) -> float:
    from sentence_transformers import util
    alvo = secao["heading"] + " " + secao["texto"][:300]
    emb = model.encode(alvo, convert_to_tensor=True)
    return float(util.cos_sim(emb, embeddings).max())


def rotear(secao: dict, termos_dict: dict, embeddings_dict: dict, model) -> tuple:
    """
    Retorna (topico | None, score, metodo).
    - topico = None → quarentena ou skip
    - metodo = 'fuzzy' | 'embed' | 'ambiguo' | 'skip'
    """
    scores = {slug: score_fuzzy(secao, termos) for slug, termos in termos_dict.items()}

    ordenados = sorted(scores.items(), key=lambda x: x[1], reverse=True)
    melhor_slug, melhor_score = ordenados[0]
    segundo_score = ordenados[1][1] if len(ordenados) > 1 else 0
    margem = melhor_score - segundo_score

    if melhor_score >= FUZZY_CONFIANTE:
        if margem >= MARGEM_EXCLUSIVA:
            return melhor_slug, melhor_score, "fuzzy"
        # Dois tópicos empatados: tenta desempatar com embedding
        if model:
            top_slugs = [s for s, v in ordenados if v >= FUZZY_CONFIANTE]
            best_embed_slug, best_embed_score = None, -1.0
            for slug in top_slugs:
                sc = score_semantico(secao, embeddings_dict[slug], model)
                if sc > best_embed_score:
                    best_embed_score, best_embed_slug = sc, slug
            if best_embed_score >= EMBED_CONFIANTE:
                return best_embed_slug, best_embed_score, f"embed-desempate"
        return None, melhor_score, "ambiguo"

    if melhor_score >= FUZZY_INCERTO and model:
        sc_e = score_semantico(secao, embeddings_dict[melhor_slug], model)
        if sc_e >= EMBED_CONFIANTE:
            return melhor_slug, sc_e, "embed"
        return None, sc_e, "embed-fraco"

    return None, melhor_score, "skip"

# ─── Helpers: escrita ─────────────────────────────────────────────────────────

def raw_path_para(slug: str, config: dict) -> Path:
    bloco = config[slug]["bloco"]
    sub   = config[slug]["sub"]
    return RAIZ_COFRE / "blocos" / bloco / sub / f"{sub}.RAW.md"


def garantir_raw(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    if not path.exists():
        path.write_text(
            f"# {path.stem} — Evidência Bruta\n\n"
            "<!-- APPEND-ONLY — gerado pelo minerador -->\n",
            encoding="utf-8"
        )


def append_raw(path: Path, secao: dict, arq: dict, score: float, metodo: str):
    ts = datetime.now().strftime("%Y-%m-%d")
    bloco = (
        f"\n---\n"
        f"**Origem:** [{arq['projeto']}] {arq['ano']} — `{arq['caminho'].name}`  \n"
        f"**Heading:** {secao['heading']}  \n"
        f"**Score:** {score:.1f} ({metodo}) | **Data:** {ts}\n\n"
        f"{secao['texto']}\n"
    )
    with path.open("a", encoding="utf-8") as f:
        f.write(bloco)


def append_inbox(secao: dict, arq: dict, score: float, motivo: str):
    ts = datetime.now().strftime("%Y-%m-%d")
    trecho = secao["texto"][:200].replace("\n", " ")
    linha = (
        f"\n[{ts}] [{arq['projeto']}] `{arq['caminho'].name}`  \n"
        f"Heading: *{secao['heading']}* | score: {score:.2f} | motivo: {motivo}  \n"
        f"> {trecho}...\n"
    )
    with INBOX_PATH.open("a", encoding="utf-8") as f:
        f.write(linha)

# ─── Manifesto ────────────────────────────────────────────────────────────────

def carregar_manifesto() -> dict:
    if MANIFESTO_PATH.exists():
        return json.loads(MANIFESTO_PATH.read_text(encoding="utf-8"))
    return {"versao": "2.0", "documentos_processados": []}


def salvar_manifesto(m: dict):
    MANIFESTO_PATH.write_text(
        json.dumps(m, ensure_ascii=False, indent=2), encoding="utf-8"
    )

# ─── Reset ────────────────────────────────────────────────────────────────────

def reset_cofre(config: dict):
    print("Resetando RAWs e manifesto...")
    for slug, info in config.items():
        path = raw_path_para(slug, config)
        if path.exists():
            path.unlink()
            print(f"  apagado: {path.name}")
    if MANIFESTO_PATH.exists():
        MANIFESTO_PATH.unlink()
        print("  apagado: manifesto.json")
    if INBOX_PATH.exists():
        # preserva o cabeçalho
        INBOX_PATH.write_text(
            "# _INBOX_REVISAR.md — Quarentena\n\n"
            "> Matches que não couberam em nenhum tópico com confiança suficiente.\n\n"
            "<!-- O minerador faz append aqui automaticamente. -->\n",
            encoding="utf-8"
        )
        print("  resetado: _INBOX_REVISAR.md")
    print("Reset completo.\n")

# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Minerador WCS — router multi-topico")
    parser.add_argument("--dry-run",        action="store_true")
    parser.add_argument("--sem-embed",      action="store_true")
    parser.add_argument("--debug-estilos",  action="store_true")
    parser.add_argument("--reset",          action="store_true")
    parser.add_argument("--sem-rede",       action="store_true", help="Não varre a rede da empresa (R:\\)")
    args = parser.parse_args()

    if args.debug_estilos:
        debug_estilos(RAIZ_SHAREPOINT)
        return

    # Carrega config de tópicos
    config = json.loads(TOPICOS_JSON.read_text(encoding="utf-8"))["topicos"]

    if args.reset:
        reset_cofre(config)

    # Carrega vocabulários
    termos_dict = {}
    for slug in config:
        vpath = RAIZ_COFRE / "vocabulario" / f"{slug}.txt"
        if not vpath.exists():
            print(f"[WARN] Vocabulário ausente, tópico ignorado: {slug}")
            continue
        termos = carregar_vocabulario(vpath)
        if termos:
            termos_dict[slug] = termos

    print(f"Tópicos ativos: {len(termos_dict)}")
    for slug, termos in termos_dict.items():
        print(f"  {slug}: {len(termos)} termos")
    print()

    # Garante que todos os diretórios e RAWs existem
    for slug in termos_dict:
        garantir_raw(raw_path_para(slug, config))

    # Descobre arquivos
    raiz_rede = None if args.sem_rede else RAIZ_REDE
    print("Varrendo SharePoint local + Rede da empresa...")
    arquivos = descobrir_docx(RAIZ_SHAREPOINT, raiz_rede)
    por_fonte = {}
    for a in arquivos:
        por_fonte[a["fonte"]] = por_fonte.get(a["fonte"], 0) + 1
    for fonte, qtd in sorted(por_fonte.items()):
        print(f"  {fonte}: {qtd} arquivo(s)")
    print(f"Total: {len(arquivos)}\n")

    if args.dry_run:
        for a in arquivos:
            print(f"  [{a['ano']}] {a['projeto']:12s}  {a['caminho'].name}")
        print(f"\nTotal: {len(arquivos)} arquivo(s) — dry-run, nada gravado.")
        return

    # Carrega embeddings (após dry-run check)
    model = None
    embeddings_dict = {}
    if not args.sem_embed:
        print("Carregando modelo de embedding...")
        from sentence_transformers import SentenceTransformer
        model = SentenceTransformer("paraphrase-multilingual-MiniLM-L12-v2")
        for slug, termos in termos_dict.items():
            embeddings_dict[slug] = model.encode(termos, convert_to_tensor=True)
        print(f"Embeddings prontos para {len(embeddings_dict)} tópicos.\n")

    # Manifesto / retomada
    manifesto = carregar_manifesto()
    processados = set(manifesto.get("documentos_processados", []))
    pendentes = [a for a in arquivos if str(a["caminho"]) not in processados]
    print(f"Ja processados: {len(processados)}  |  Pendentes: {len(pendentes)}\n")

    # Contadores por tópico
    contadores = {slug: {"raw": 0, "inbox": 0, "skip": 0} for slug in termos_dict}
    n_erros = 0

    for arq in tqdm(pendentes, desc="Minerando", unit="doc"):
        try:
            secoes = extrair_secoes(arq["caminho"])
        except Exception as e:
            tqdm.write(f"[WARN] {arq['caminho'].name}: {e}")
            n_erros += 1
            continue

        for secao in secoes:
            if not secao["texto"].strip():
                continue

            topico, score, metodo = rotear(secao, termos_dict, embeddings_dict, model)

            if topico:
                append_raw(raw_path_para(topico, config), secao, arq, score, metodo)
                contadores[topico]["raw"] += 1
            elif metodo in ("ambiguo", "embed-fraco"):
                append_inbox(secao, arq, score, metodo)
                for slug in contadores:
                    contadores[slug]["inbox"] += 0  # conta no total abaixo
            else:
                pass  # skip silencioso

        processados.add(str(arq["caminho"]))

    # Persiste manifesto
    manifesto["documentos_processados"] = sorted(processados)
    manifesto["ultima_mineracao"] = datetime.now().isoformat()
    salvar_manifesto(manifesto)

    # Relatório final
    print("\n--- Resultado por topico ---")
    total_raw = 0
    for slug, c in contadores.items():
        if c["raw"] > 0:
            print(f"  {slug:20s}: {c['raw']:4d} secoes -> RAW")
            total_raw += c["raw"]
    print(f"\nTotal RAW    : {total_raw}")
    print(f"Erros leitura: {n_erros}")
    print(f"(Quarentena em _INBOX_REVISAR.md)")


if __name__ == "__main__":
    main()
